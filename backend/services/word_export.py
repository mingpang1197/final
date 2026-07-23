from __future__ import annotations

"""이지리드 문서 Word(.docx) 내보내기.

역할: DocumentResponse를 Easy-Read 타이포그래피 규칙에 맞는 Word 파일로 변환한다.
주요 기능: export_to_docx — 소제목 + 항목별 글상자(투명 테두리) 2단(왼쪽 그림 / 오른쪽 본문).
관계: export_layout, image_assets, models/schemas, routers/documents(export API).
"""

import base64
import io
import re
import tempfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from backend.models.schemas import DocumentResponse, ImagePlacement
from backend.services.export_layout import (
    align_placements_one_per_section,
    is_image_placeholder,
    parse_export_sections,
    parse_section_items,
    prepare_placements_for_export,
    split_item_lines_into_blocks,
)
from backend.services.easy_read_sanitize import split_standard_closing
from backend.services.judgment_merge import (
    EASY_READ_PROVISION_PARAGRAPHS,
    split_judgment_at_reason,
)
from backend.services.image_assets import resolve_placement_image
from backend.services.image_matcher import MAX_IMAGES_PER_TEXT, find_images_for_line
from backend.services.rich_text import has_style_markers, iter_styled_runs
from backend.services.court_fonts import (
    ExportFontProfile,
    bundled_court_font_profile,
    enable_word_font_embedding,
)

BODY_PT = 12
BOLD_BODY_PT = 12
HEADING_PT = 17
FOOTER_PT = 11
LINE_SPACING = 2.0  # 200% (이지리드 가이드)
PAGE_MARGIN_IN = 0.6
PAGE_MARGIN = Inches(PAGE_MARGIN_IN)
# Word A4 본문 높이(여백 제외) — 소제목 페이지 중단 이후면 다음 페이지 추정용
EASY_READ_USABLE_PAGE_PT = 680.0
EASY_READ_PROVISION_EST_PT = 200.0
EASY_READ_HEADING_EST_PT = 56.0
EASY_READ_ITEM_BASE_EST_PT = 118.0
EASY_READ_LINE_EST_PT = 26.0
IMAGE_COL_IN = 2.05
GAP_COL_IN = 0.1
BODY_COL_IN = 8.5 - PAGE_MARGIN_IN * 2 - IMAGE_COL_IN - GAP_COL_IN
IMAGE_COL_WIDTH = Inches(IMAGE_COL_IN)
GAP_COL_WIDTH = Inches(GAP_COL_IN)
BODY_COL_WIDTH = Inches(BODY_COL_IN)
IMAGE_DISPLAY_WIDTH = Inches(1.85)
IMAGE_CELL_FILL = "FFFFFF"
CELL_PAD_V = 40  # twips — 셀 상하 여백
CELL_PAD_GAP = 48  # twips — 그림·본문 사이 간격

_SKIP_LINE = re.compile(r"^(---+\s*|\(\d+/\d+\s*쪽\)|\d+/\d+\s*쪽|>\s)")
_META_SECTION_START = re.compile(r"^###\s*수정\s*사항")

DEFAULT_ASCII_FONT = "Malgun Gothic"
DEFAULT_EAST_ASIA_FONT = "맑은 고딕"
EASY_READ_COURT_FONT = "CourtBT"
_GOTHIC_FONTS = frozenset(
    {"malgun gothic", "맑은 고딕", "nanumgothic", "나눔고딕", "dotum", "돋움", "gulim", "굴림"}
)

EASY_READ_FONT_PROFILE = bundled_court_font_profile()

_FONT_PROFILE_STACK: list[ExportFontProfile] = []


def _active_font_profile() -> ExportFontProfile:
    if _FONT_PROFILE_STACK:
        return _FONT_PROFILE_STACK[-1]
    return ExportFontProfile()


def _push_font_profile(profile: ExportFontProfile | None) -> None:
    _FONT_PROFILE_STACK.append(profile or ExportFontProfile())


def _pop_font_profile() -> None:
    if _FONT_PROFILE_STACK:
        _FONT_PROFILE_STACK.pop()


def _iter_all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _normalized_reason(text: str) -> str:
    return re.sub(r"\s+", "", (text or "").strip())


def is_reason_heading_paragraph(paragraph) -> bool:
    """문단 첫 줄만 「이유」인지 (pdf2docx·OCR 공통)."""
    raw = (paragraph.text or "").replace("\r", "")
    combined = "".join(run.text for run in paragraph.runs)
    for candidate in (raw, combined):
        if not candidate.strip():
            continue
        first_line = candidate.split("\n", 1)[0].strip()
        if _normalized_reason(first_line) == "이유":
            return True
    return _normalized_reason(raw.split("\n", 1)[0].strip()) == "이유"


def _paragraphs_in_table_element(tbl_el, doc: Document):
    from docx.text.paragraph import Paragraph

    for row in tbl_el.findall(qn("w:tr")):
        for cell in row.findall(qn("w:tc")):
            for el in cell:
                if el.tag == qn("w:p"):
                    yield Paragraph(el, doc)
                elif el.tag == qn("w:tbl"):
                    yield from _paragraphs_in_table_element(el, doc)


def iter_body_paragraphs_in_order(doc: Document):
    """본문 XML 순서대로 단락 (표 안 포함)."""
    from docx.text.paragraph import Paragraph

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield from _paragraphs_in_table_element(child, doc)


def find_reason_anchor_paragraph(doc: Document):
    """문서 흐름상 「이유」 제목 단락 — pdf2docx 순서 오류 대비 마지막 매칭."""
    found = None
    for paragraph in iter_body_paragraphs_in_order(doc):
        if is_reason_heading_paragraph(paragraph):
            found = paragraph
    return found


def _split_reason_heading_paragraph(paragraph):
    """「이유」와 본문이 한 단락에 붙어 있으면 제목만 남기고 나머지는 다음 단락으로."""
    from docx.text.paragraph import Paragraph

    raw = (paragraph.text or "").replace("\r", "")
    lines = raw.split("\n")
    if len(lines) <= 1:
        return paragraph
    if _normalized_reason(lines[0].strip()) != "이유":
        return paragraph

    p_el = paragraph._element
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)
    paragraph.add_run(lines[0].strip())

    rest = "\n".join(lines[1:]).strip()
    if not rest:
        return paragraph

    new_p = OxmlElement("w:p")
    p_el.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(rest)
    return paragraph


def prepare_reason_insert_anchor(doc: Document):
    """삽입 직전 「이유」 단락 — 분리·문서 순서 기준."""
    anchor = find_reason_anchor_paragraph(doc)
    if anchor is None:
        return None
    return _split_reason_heading_paragraph(anchor)


def _font_name_key(name: str) -> str:
    return (name or "").strip().lower()


def _pick_dominant_font(counter: Counter[str]) -> str | None:
    if not counter:
        return None
    ranked = counter.most_common()
    best_name, _ = ranked[0]
    for name, weight in ranked:
        key = _font_name_key(name)
        if key in _GOTHIC_FONTS:
            continue
        if any(hint in key for hint in ("batang", "myung", "명조", "바탕", "gungsuh", "궁서", "hy")):
            return name
    for name, _ in ranked:
        if _font_name_key(name) not in _GOTHIC_FONTS:
            return name
    return best_name


def _collect_run_fonts(run, weight: int, east_asia: Counter[str], ascii_fonts: Counter[str], sizes: Counter[float]) -> None:
    r_pr = run._element.find(qn("w:rPr"))
    east_name: str | None = None
    ascii_name: str | None = None
    if r_pr is not None:
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is not None:
            east_name = r_fonts.get(qn("w:eastAsia"))
            ascii_name = r_fonts.get(qn("w:ascii")) or r_fonts.get(qn("w:hAnsi"))
        sz = r_pr.find(qn("w:sz"))
        if sz is not None and sz.get(qn("w:val")):
            try:
                sizes[int(sz.get(qn("w:val"))) / 2] += weight
            except (TypeError, ValueError):
                pass
    if not east_name and run.font.name:
        east_name = run.font.name
    if east_name:
        east_asia[east_name] += weight
    if ascii_name:
        ascii_fonts[ascii_name] += weight
    elif east_name:
        ascii_fonts[east_name] += weight


def infer_font_profile_from_reason_vicinity(doc: Document, anchor_paragraph) -> ExportFontProfile:
    """「이유」 바로 다음 본문 구간 글꼴을 이지리드에 맞춘다."""
    east_asia: Counter[str] = Counter()
    ascii_fonts: Counter[str] = Counter()
    sizes: Counter[float] = Counter()
    passed_anchor = False
    collected = 0

    for paragraph in iter_body_paragraphs_in_order(doc):
        if paragraph._element is anchor_paragraph._element:
            passed_anchor = True
            continue
        if not passed_anchor:
            continue
        if is_reason_heading_paragraph(paragraph):
            break
        text = paragraph.text or ""
        if not text.strip():
            continue
        for run in paragraph.runs:
            chunk = run.text or ""
            if not chunk.strip():
                continue
            _collect_run_fonts(run, len(chunk), east_asia, ascii_fonts, sizes)
        collected += len(text)
        if collected >= 4000:
            break

    east = _pick_dominant_font(east_asia)
    ascii_dom = _pick_dominant_font(ascii_fonts) or east
    if east:
        body_pt = sizes.most_common(1)[0][0] if sizes else BODY_PT
        return ExportFontProfile(
            ascii=ascii_dom or east,
            h_ansi=ascii_dom or east,
            east_asia=east,
            body_pt=body_pt,
        )
    return infer_font_profile_from_document(doc)


def infer_font_profile_from_document(doc: Document) -> ExportFontProfile:
    """pdf2docx 원문 DOCX에서 가장 많이 쓰인 글꼴·본문 크기 추정."""
    font_weights: Counter[str] = Counter()
    size_weights: Counter[float] = Counter()

    for paragraph in _iter_all_paragraphs(doc):
        for run in paragraph.runs:
            text = run.text or ""
            if not text.strip():
                continue
            weight = len(text)
            r_pr = run._element.find(qn("w:rPr"))
            if r_pr is not None:
                r_fonts = r_pr.find(qn("w:rFonts"))
                if r_fonts is not None:
                    for key in ("eastAsia", "ascii", "hAnsi"):
                        name = r_fonts.get(qn(f"w:{key}"))
                        if name:
                            font_weights[name] += weight
                            break
                else:
                    if run.font.name:
                        font_weights[run.font.name] += weight
                sz = r_pr.find(qn("w:sz"))
                if sz is not None and sz.get(qn("w:val")):
                    try:
                        size_weights[int(sz.get(qn("w:val"))) / 2] += weight
                    except (TypeError, ValueError):
                        pass
            elif run.font.name:
                font_weights[run.font.name] += weight

    if not font_weights:
        return ExportFontProfile(
            ascii=DEFAULT_ASCII_FONT,
            h_ansi=DEFAULT_ASCII_FONT,
            east_asia=DEFAULT_EAST_ASIA_FONT,
            body_pt=BODY_PT,
        )

    dominant = font_weights.most_common(1)[0][0]
    body_pt = size_weights.most_common(1)[0][0] if size_weights else BODY_PT
    return ExportFontProfile(
        ascii=dominant,
        h_ansi=dominant,
        east_asia=dominant,
        body_pt=body_pt,
    )


def _body_pt() -> float:
    return _active_font_profile().body_pt


def _set_run_font(run, size_pt: float, bold: bool = False) -> None:
    profile = _active_font_profile()
    ascii_name = profile.ascii
    east_name = profile.east_asia
    run.font.name = ascii_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_name)
    r_fonts.set(qn("w:hAnsi"), profile.h_ansi)
    r_fonts.set(qn("w:eastAsia"), east_name)
    for tag in ("w:b", "w:bCs"):
        existing = r_pr.find(qn(tag))
        if bold:
            if existing is None:
                r_pr.append(OxmlElement(tag))
        elif existing is not None:
            r_pr.remove(existing)


def _is_heading(line: str) -> bool:
    s = line.strip()
    if s.startswith("<") and s.endswith(">"):
        return True
    if s.startswith("■"):
        return True
    return s.startswith("#")


def _clean_heading(line: str) -> str:
    return re.sub(r"^#+\s*", "", line.strip())


def _configure_section(section) -> None:
    section.top_margin = PAGE_MARGIN
    section.bottom_margin = PAGE_MARGIN
    section.left_margin = PAGE_MARGIN
    section.right_margin = PAGE_MARGIN
    _set_page_number_footer(section)


def _inches_to_twips(width: Inches) -> int:
    return int(width.inches * 1440)


def _set_page_number_footer(section) -> None:
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()

    run_prefix = fp.add_run("- ")
    _set_run_font(run_prefix, FOOTER_PT)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_prefix._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run_prefix._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_prefix._r.append(fld_sep)

    run_num = fp.add_run("1")
    _set_run_font(run_num, FOOTER_PT)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_num._r.append(fld_end)

    run_suffix = fp.add_run(" -")
    _set_run_font(run_suffix, FOOTER_PT)


def _remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def _set_row_cant_split(row) -> None:
    """행이 페이지 경계에서 쪼개지지 않도록 — 소제목+본문 블록을 다음 페이지로."""
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    for child in list(tr_pr):
        if child.tag == qn("w:cantSplit"):
            tr_pr.remove(child)
    tr_pr.append(OxmlElement("w:cantSplit"))


def _open_keep_together_cell(doc: Document):
    """레거시 — bordered table 행 방식으로 대체됨."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _remove_table_borders(table)
    _set_row_cant_split(table.rows[0])
    cell = table.rows[0].cells[0]
    _set_cell_borders_transparent(cell)
    _clear_cell_shading(cell)
    _reset_cell_paragraphs(cell)
    return cell


def _estimate_easy_read_section_height(section, by_item: dict) -> float:
    height = EASY_READ_HEADING_EST_PT if section.heading else 0.0
    for item in parse_section_items(section):
        line_count = len([ln for ln in item.lines if ln.strip()])
        height += EASY_READ_ITEM_BASE_EST_PT + line_count * EASY_READ_LINE_EST_PT
    return height


def _add_page_break_before_cell_content(cell) -> None:
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _export_easy_read_layout_in_table(
    table,
    text: str,
    placements: list[ImagePlacement],
    *,
    page_fill_pt: float = 0.0,
) -> None:
    """소제목마다 cantSplit 행 + (추정) 페이지 절반 이후면 소제목을 다음 페이지부터."""
    from backend.services.word_textbox import append_easy_read_section_row

    body, closing = split_standard_closing(text)
    by_item = align_placements_one_per_section(body, placements)
    fill = page_fill_pt

    for section_index, section in enumerate(parse_export_sections(body)):
        section_est = _estimate_easy_read_section_height(section, by_item)
        pos_on_page = fill % EASY_READ_USABLE_PAGE_PT
        need_break = section_index > 0 and pos_on_page >= EASY_READ_USABLE_PAGE_PT / 2

        cell = append_easy_read_section_row(table)
        _reset_cell_paragraphs(cell)
        if need_break:
            _add_page_break_before_cell_content(cell)
            fill = (fill // EASY_READ_USABLE_PAGE_PT + 1) * EASY_READ_USABLE_PAGE_PT

        if section.heading:
            _add_heading_paragraph(cell, section.heading)
        for item in parse_section_items(section):
            placement = by_item.get(item.start_line_index)
            if placement is not None and not isinstance(placement, ImagePlacement):
                placement = ImagePlacement(**placement)  # type: ignore[arg-type]
            if placement:
                blocks = split_item_lines_into_blocks(item.lines)
                _add_item_text_boxes(cell, placement, blocks[0])
                for block_lines in blocks[1:]:
                    _add_item_text_boxes(cell, None, block_lines)
            else:
                _add_item_text_boxes(cell, None, item.lines)
        fill += section_est

    if closing:
        if fill % EASY_READ_USABLE_PAGE_PT >= EASY_READ_USABLE_PAGE_PT / 2:
            cell = append_easy_read_section_row(table)
            _reset_cell_paragraphs(cell)
            _add_page_break_before_cell_content(cell)
        else:
            cell = append_easy_read_section_row(table)
            _reset_cell_paragraphs(cell)
        _add_closing_paragraph(cell, closing)


def _set_cell_borders_transparent(cell) -> None:
    """글상자처럼 셀 테두리를 투명(nil) 처리."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tc_pr.append(borders)


def _set_cell_vertical_align(cell, *, top: bool = True) -> None:
    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.TOP if top else WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )


def _set_cell_margins(cell, *, top: int = 0, bottom: int = 0, left: int = 0, right: int = 0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for edge, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _reset_cell_paragraphs(cell) -> None:
    while len(cell.paragraphs) > 1:
        cell._element.remove(cell.paragraphs[-1]._element)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shading)


def _clear_cell_shading(cell) -> None:
    """항목 내부 2단 표 — 배경·글상자 느낌 제거."""
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:shd"):
            tc_pr.remove(child)


def _set_column_widths(table, widths: list[Inches]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def _set_cell_width_fixed(cell, width: Inches) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:tcW"):
            tc_pr.remove(child)
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(_inches_to_twips(width)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.insert(0, tc_w)


def _set_table_fixed_layout(table, widths: list[Inches]) -> None:
    """행마다 열 너비가 흔들리지 않도록 고정(그림 유무 본문 정렬 일치)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    for child in list(tbl_pr):
        if child.tag == qn("w:tblLayout"):
            tbl_pr.remove(child)
        if child.tag == qn("w:tblInd"):
            tbl_pr.remove(child)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    for existing in tbl.findall(qn("w:tblGrid")):
        tbl.remove(existing)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(_inches_to_twips(width)))
        tbl_grid.append(grid_col)
    tbl_pr.addnext(tbl_grid)


def _resolve_placement_path(placement: ImagePlacement) -> Path | None:
    raw = (placement.image_base64 or "").strip()
    if raw:
        payload = raw.split(",", 1)[1] if raw.startswith("data:") else raw
        try:
            data = base64.b64decode(payload)
        except (ValueError, TypeError):
            return None
        suffix = ".png"
        lower = raw.lower()
        for ext in (".jpg", ".jpeg", ".webp", ".gif"):
            if ext in lower[:40]:
                suffix = ext
                break
        tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
        tmp.write(data)
        tmp.close()
        return Path(tmp.name)
    return resolve_placement_image(
        image_file=placement.image_file,
        image_url=placement.image_url,
    )


def _apply_body_format(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = LINE_SPACING
    paragraph.paragraph_format.space_after = Pt(6)


def _add_runs_to_paragraph(paragraph, line: str, *, size_pt: float, bold_default: bool = False) -> None:
    for part, is_bold, run_pt in iter_styled_runs(line, default_pt=size_pt):
        run = paragraph.add_run(part)
        _set_run_font(run, run_pt, bold=is_bold or bold_default)


def _add_heading_paragraph(doc_or_cell, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return
    p = doc_or_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    raw = _clean_heading(stripped) if _is_heading(stripped) else stripped
    if has_style_markers(stripped):
        _add_runs_to_paragraph(p, stripped, size_pt=HEADING_PT)
    else:
        run = p.add_run(raw)
        _set_run_font(run, HEADING_PT, bold=True)


def _add_body_paragraph_to_cell(cell, line: str, *, first: bool = False) -> None:
    stripped = line.strip()
    if not stripped or _SKIP_LINE.match(stripped):
        return
    if is_image_placeholder(stripped):
        return
    if first and cell.paragraphs and not cell.paragraphs[0].text.strip():
        p = cell.paragraphs[0]
        p.clear()
    else:
        p = cell.add_paragraph()
    _apply_body_format(p)
    if first:
        p.paragraph_format.space_before = Pt(0)
    _add_runs_to_paragraph(p, stripped, size_pt=_body_pt())


def _add_picture_to_cell(cell, placement: ImagePlacement) -> None:
    img_path = _resolve_placement_path(placement)
    if not img_path:
        return
    _reset_cell_paragraphs(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(img_path), width=IMAGE_DISPLAY_WIDTH)


def _add_item_text_boxes(
    doc_or_cell,
    placement: ImagePlacement | None,
    body_lines: list[str],
) -> None:
    """항목 1개 = (그림 | 본문) 2단. 그림 없으면 왼쪽 빈칸 유지(그림 탭과 동일)."""
    table = doc_or_cell.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    _remove_table_borders(table)
    col_widths = [IMAGE_COL_WIDTH, Inches(BODY_COL_IN + GAP_COL_IN)]
    _set_table_fixed_layout(table, col_widths)
    _set_column_widths(table, col_widths)

    row = table.rows[0]
    _set_row_cant_split(row)
    row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
    image_cell = row.cells[0]
    body_cell = row.cells[1]
    _set_cell_width_fixed(image_cell, col_widths[0])
    _set_cell_width_fixed(body_cell, col_widths[1])

    for cell in (image_cell, body_cell):
        _set_cell_borders_transparent(cell)
        _clear_cell_shading(cell)
        _set_cell_vertical_align(cell, top=True)
        _reset_cell_paragraphs(cell)

    if placement:
        _add_picture_to_cell(image_cell, placement)
    else:
        p = image_cell.paragraphs[0]
        run = p.add_run("\u00a0")
        _set_run_font(run, _body_pt())
    _set_cell_margins(image_cell, top=CELL_PAD_V, bottom=CELL_PAD_V, left=0, right=CELL_PAD_GAP)
    _set_cell_margins(body_cell, top=CELL_PAD_V, bottom=CELL_PAD_V, left=CELL_PAD_GAP, right=0)

    first = True
    for line in body_lines:
        _add_body_paragraph_to_cell(body_cell, line, first=first)
        first = False

    spacer = doc_or_cell.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(12)
    spacer.paragraph_format.keep_with_next = True


def _add_closing_paragraph(doc_or_cell, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return
    p = doc_or_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    _apply_body_format(p)
    _add_runs_to_paragraph(p, stripped, size_pt=_body_pt())


def _export_easy_read_layout(
    doc: Document,
    text: str,
    placements: list[ImagePlacement],
) -> None:
    """비표 레이아웃(레거시). 내보내기는 _export_easy_read_layout_in_table 사용."""
    body, closing = split_standard_closing(text)
    by_item = align_placements_one_per_section(body, placements)

    for section in parse_export_sections(body):
        block = _open_keep_together_cell(doc)
        if section.heading:
            _add_heading_paragraph(block, section.heading)
        for item in parse_section_items(section):
            placement = by_item.get(item.start_line_index)
            if placement is not None and not isinstance(placement, ImagePlacement):
                placement = ImagePlacement(**placement)  # type: ignore[arg-type]
            if placement:
                blocks = split_item_lines_into_blocks(item.lines)
                _add_item_text_boxes(block, placement, blocks[0])
                for block_lines in blocks[1:]:
                    _add_item_text_boxes(block, None, block_lines)
            else:
                _add_item_text_boxes(block, None, item.lines)

    if closing:
        _add_closing_paragraph(doc, closing)


def _add_rich_paragraph(doc_or_cell, line: str) -> None:
    stripped = line.strip()
    if not stripped or _SKIP_LINE.match(stripped):
        return
    if stripped == "## 수정된 이지리드 번역본":
        return
    if is_image_placeholder(stripped):
        return

    if _is_heading(stripped):
        _add_heading_paragraph(doc_or_cell, stripped)
        return

    p = doc_or_cell.add_paragraph()
    _apply_body_format(p)
    _add_runs_to_paragraph(p, stripped, size_pt=_body_pt())


def _add_picture(doc: Document, image_file: str, image_url: str | None = None) -> None:
    img_path = resolve_placement_image(image_file=image_file, image_url=image_url)
    if not img_path:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(str(img_path), width=IMAGE_DISPLAY_WIDTH)


def _export_text_to_cell(cell, text: str, *, skip_meta: bool = True) -> None:
    in_meta_section = False
    inserted_images: set[str] = set()

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if skip_meta and stripped.startswith("원본 파일:"):
            continue
        if skip_meta and stripped == "## 수정된 이지리드 번역본":
            continue
        if skip_meta and _META_SECTION_START.match(stripped):
            in_meta_section = True
            continue
        if in_meta_section:
            continue

        _add_rich_paragraph(cell, line)

        if len(inserted_images) >= MAX_IMAGES_PER_TEXT:
            continue
        for match in find_images_for_line(
            stripped,
            exclude=inserted_images,
            max_total=MAX_IMAGES_PER_TEXT,
        ):
            img_path = resolve_placement_image(image_file=match.image_file)
            if img_path:
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run()
                run.add_picture(str(img_path), width=IMAGE_DISPLAY_WIDTH)
            inserted_images.add(match.image_file)


def _export_text(doc: Document, text: str, *, skip_meta: bool = True) -> None:
    in_meta_section = False
    inserted_images: set[str] = set()

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if skip_meta and stripped.startswith("원본 파일:"):
            continue
        if skip_meta and stripped == "## 수정된 이지리드 번역본":
            continue
        if skip_meta and _META_SECTION_START.match(stripped):
            in_meta_section = True
            continue
        if in_meta_section:
            continue

        _add_rich_paragraph(doc, line)

        if len(inserted_images) >= MAX_IMAGES_PER_TEXT:
            continue
        for match in find_images_for_line(
            stripped,
            exclude=inserted_images,
            max_total=MAX_IMAGES_PER_TEXT,
        ):
            _add_picture(doc, match.image_file)
            inserted_images.add(match.image_file)


def _export_judgment_plain(doc: Document, text: str) -> None:
    """OCR 원문 판결 — 줄바꿈 유지, 자동 그림 매칭 없음."""
    for line in text.split("\n"):
        if not line.strip():
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        _apply_body_format(p)
        run = p.add_run(line.rstrip())
        _set_run_font(run, _body_pt())


def _export_easy_read_provision(doc_or_cell) -> None:
    """이유 직후 — Easy-Read 제공 고지(가·나)."""
    for block in EASY_READ_PROVISION_PARAGRAPHS:
        lines = block.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("※"):
                p = doc_or_cell.add_paragraph()
                _apply_body_format(p)
                run = p.add_run(stripped)
                _set_run_font(run, _body_pt(), bold=True)
            elif stripped.startswith(("가.", "나.")):
                p = doc_or_cell.add_paragraph()
                _apply_body_format(p)
                run = p.add_run(stripped)
                _set_run_font(run, _body_pt(), bold=True)
            else:
                _add_rich_paragraph(doc_or_cell, line)
        if len(lines) > 1:
            spacer = doc_or_cell.add_paragraph()
            spacer.paragraph_format.space_after = Pt(4)


def _export_easy_read_body(doc: Document, easy_body: str, raw_placements: list[ImagePlacement]) -> None:
    placements = prepare_placements_for_export(easy_body, raw_placements)
    sections = parse_export_sections(easy_body)
    if any(section.heading for section in sections):
        _export_easy_read_layout(doc, easy_body, placements)
    else:
        _export_text(doc, easy_body, skip_meta=True)


def _collect_placements(doc: DocumentResponse) -> list[ImagePlacement] | None:
    for segment in doc.translation_segments or []:
        if segment.image_placements:
            return segment.image_placements
    return None


def _collect_body_text(doc: DocumentResponse) -> str:
    if doc.translation_segments:
        for segment in doc.translation_segments:
            if segment.image_placements and segment.easy_text and segment.easy_text.strip():
                return segment.easy_text.strip()
    if doc.translation_text and doc.translation_text.strip():
        return doc.translation_text.strip()
    if doc.translation_segments:
        parts = [
            s.easy_text.strip()
            for s in doc.translation_segments
            if s.easy_text and s.easy_text.strip()
        ]
        if parts:
            return max(parts, key=len)
    if doc.summary and doc.summary.strip():
        return doc.summary.strip()
    return ""


def collect_body_text(doc: DocumentResponse) -> str:
    return _collect_body_text(doc)


def _write_easy_read_bordered(
    host: Document,
    doc: DocumentResponse,
    *,
    font_profile: ExportFontProfile | None = None,
) -> None:
    easy_body = _collect_body_text(doc)
    if not easy_body:
        return

    from backend.services.word_textbox import (
        append_easy_read_section_row,
        begin_easy_read_outer_table,
        finish_easy_read_outer_table,
    )

    _push_font_profile(font_profile if font_profile is not None else EASY_READ_FONT_PROFILE)
    try:
        table, cell = begin_easy_read_outer_table(host)
        _reset_cell_paragraphs(cell)
        _export_easy_read_provision(cell)
        raw_placements = _collect_placements(doc) or []
        placements = prepare_placements_for_export(easy_body, raw_placements)
        sections = parse_export_sections(easy_body)
        if any(section.heading for section in sections):
            _export_easy_read_layout_in_table(
                table,
                easy_body,
                placements,
                page_fill_pt=EASY_READ_PROVISION_EST_PT,
            )
        else:
            body_cell = append_easy_read_section_row(table)
            _reset_cell_paragraphs(body_cell)
            _export_text_to_cell(body_cell, easy_body, skip_meta=True)
    finally:
        _pop_font_profile()
    finish_easy_read_outer_table(host)


def _append_easy_read_in_textbox(host: Document, doc: DocumentResponse) -> None:
    _write_easy_read_bordered(host, doc)


def build_easy_read_insert_document(
    doc: DocumentResponse,
    *,
    font_profile: ExportFontProfile | None = None,
) -> Document:
    """「이유」 직후 삽입용 — 고지·이지리드 본문을 테두리 표(행별 cantSplit)에 담는다."""
    easy_body = _collect_body_text(doc)
    if not easy_body:
        return Document()

    host = Document()
    for section in host.sections:
        section.top_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN
        section.left_margin = PAGE_MARGIN
        section.right_margin = PAGE_MARGIN
        section.footer.is_linked_to_previous = False
        if section.footer.paragraphs:
            section.footer.paragraphs[0].clear()

    _write_easy_read_bordered(host, doc, font_profile=font_profile)
    page_break = host.add_paragraph()
    page_break.paragraph_format.space_before = Pt(0)
    page_break.paragraph_format.space_after = Pt(0)
    page_break.add_run().add_break(WD_BREAK.PAGE)
    enable_word_font_embedding(host)
    return host


def export_to_docx(
    doc: DocumentResponse,
    *,
    include_meta: bool = False,
    source_file: Path | None = None,
    merge_source_pdf: bool = True,
) -> bytes:
    word = Document()

    for section in word.sections:
        _configure_section(section)

    if include_meta:
        title = word.add_heading("이지리드 판결문", level=0)
        for run in title.runs:
            _set_run_font(run, 18, bold=True)
        meta = word.add_paragraph(f"원본 파일: {doc.filename}  |  유형: {doc.doc_type}")
        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT

    easy_body = _collect_body_text(doc)
    if not easy_body:
        buffer = io.BytesIO()
        word.save(buffer)
        return buffer.getvalue()

    if source_file is not None and merge_source_pdf:
        from backend.services.pdf_source_merge import merge_pdf_with_easy_read_insert

        merged = merge_pdf_with_easy_read_insert(source_file, doc)
        if merged:
            return merged

    split = split_judgment_at_reason(doc.full_text or "") if (doc.full_text or "").strip() else None

    if split and easy_body:
        prefix, suffix = split
        _export_judgment_plain(word, prefix)
        _append_easy_read_in_textbox(word, doc)
        if suffix.strip():
            word.add_page_break()
            _export_judgment_plain(word, suffix)
    elif easy_body:
        _append_easy_read_in_textbox(word, doc)

    enable_word_font_embedding(word)
    buffer = io.BytesIO()
    word.save(buffer)
    return buffer.getvalue()
