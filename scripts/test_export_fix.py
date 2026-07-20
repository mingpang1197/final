"""Verify word export fixes."""
import io
import json
import sqlite3
import sys

from docx import Document

from backend.models.schemas import DocumentResponse, TranslationSegment
from backend.services.image_matcher import find_matching_images
from backend.services.word_export import export_to_docx


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")

    meta_sample = """## 수정된 이지리드 번역본
## <김OO 가 원하는 것>
김OO은 **지적장애**로 장애인 등록을 신청했습니다.

### 수정 사항 설명
- should be removed
"""
    doc = DocumentResponse(
        id="t",
        filename="test.pdf",
        doc_type="administrative",
        stage="translated",
        page_count=1,
        full_text="",
        translation_segments=[
            TranslationSegment(id="1", original="x", easy_text=meta_sample, source="solar")
        ],
    )
    texts = [
        p.text.strip()
        for p in Document(io.BytesIO(export_to_docx(doc))).paragraphs
        if p.text.strip()
    ]
    assert any("김OO" in t for t in texts), "body missing"
    assert not any("should be removed" in t for t in texts), "meta not removed"
    print("Test1 meta bug: OK")

    spam_segments = [
        TranslationSegment(
            id="1",
            original="x",
            easy_text="김OO은 지적장애로 장애인 등록을 신청했습니다.",
            source="solar",
        ),
        TranslationSegment(
            id="2", original="", easy_text="", title="소송비용\n(원고 10%)", source="db"
        ),
        TranslationSegment(id="3", original="", easy_text="", title="각하", source="db"),
    ]
    doc2 = DocumentResponse(
        id="t2",
        filename="50969.pdf",
        doc_type="administrative",
        stage="translated",
        page_count=1,
        full_text="",
        translation_segments=spam_segments,
    )
    texts2 = [
        p.text.strip()
        for p in Document(io.BytesIO(export_to_docx(doc2))).paragraphs
        if p.text.strip()
    ]
    assert not any(t == "각하" or t.startswith("<각하") for t in texts2)
    assert any("지적장애" in t for t in texts2)
    print("Test2 image-only segments: OK")

    sample = "김OO은 지적장애로 장애인 등록을 신청했습니다. 양천구청장은 등록을 거부했습니다."
    matches = find_matching_images(sample)
    assert len(matches) <= 8
    print(f"Test3 image matches: {len(matches)} (<=8)")

    conn = sqlite3.connect("data/app.db")
    for row in conn.execute(
        "select filename, translation_json from documents where translation_json is not null"
    ):
        if "50969" not in row[0]:
            continue
        data = json.loads(row[1])
        segs = data if isinstance(data, list) else data.get("segments", [])
        if not segs:
            continue
        doc3 = DocumentResponse(
            id="db",
            filename=row[0],
            doc_type="administrative",
            stage="translated",
            page_count=1,
            full_text="",
            translation_segments=[TranslationSegment(**s) for s in segs],
        )
        texts3 = [
            p.text.strip()
            for p in Document(io.BytesIO(export_to_docx(doc3))).paragraphs
            if p.text.strip()
        ]
        chars = sum(len(t) for t in texts3)
        print(f"Test4 DB 50969: paragraphs={len(texts3)} chars={chars}")
        assert chars > 500, f"50969 export too short: {chars}"
        break
    else:
        print("Test4 DB 50969: skipped")

    print("ALL TESTS PASSED")


if __name__ == "__main__":
    main()
